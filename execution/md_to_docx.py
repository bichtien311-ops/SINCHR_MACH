"""
Конвертер Markdown -> Word (.docx) для курса «Синхронные машины».

Возможности:
- заголовки (#..####), абзацы с **жирным** и `моноширинным` текстом;
- блоки кода ```...```;
- таблицы в формате Markdown (| ... |);
- маркированные (-) и нумерованные (1.) списки;
- цитаты (>), горизонтальные линии (---).

Запуск:
    python execution/md_to_docx.py

Скрипт:
1) конвертирует корневые документы (программа, литература, вопросы) в Результаты/;
2) конвертирует все Лекция.md (папки Л01..Л12) и лабораторные (папка Лабораторные/);
3) собирает единую РАСШИРЕННУЮ методичку (лекции + лабораторные) —
   Методичка_Синхронные_машины.docx.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Корень курса и папка результатов
ROOT = Path(__file__).resolve().parent.parent
COURSE_DIR = ROOT / "Синхронные_машины"
LAB_DIR = COURSE_DIR / "Лабораторные"
RESULT_DIR = ROOT / "Результаты"

INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
IMAGE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def add_image(doc: Document, caption: str, rel_path: str, base_dir) -> None:
    """Вставить изображение с подписью. Путь ищется относительно base_dir,
    затем относительно COURSE_DIR."""
    candidates = []
    if base_dir is not None:
        candidates.append((base_dir / rel_path).resolve())
    candidates.append((COURSE_DIR / rel_path).resolve())
    candidates.append((ROOT / rel_path).resolve())
    img = next((c for c in candidates if c.exists()), None)
    if img is None:
        p = doc.add_paragraph()
        p.add_run(f"[рисунок не найден: {rel_path}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img), width=Inches(5.3))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)


def add_runs(paragraph, text: str) -> None:
    """Добавить текст в абзац с поддержкой **жирного** и `кода`."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_code_block(doc: Document, lines: list[str]) -> None:
    """Блок кода моноширинным шрифтом."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    """Таблица из строк Markdown (первая строка — заголовок)."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            cells[j].text = ""
            add_runs(cells[j].paragraphs[0], text)
            if i == 0:
                for run in cells[j].paragraphs[0].runs:
                    run.bold = True


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1] + "|"


def is_table_separator(line: str) -> bool:
    s = line.strip().strip("|")
    return bool(s) and all(c in "-: " for c in s)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_markdown(doc: Document, md_text: str, base_dir=None) -> None:
    """Отрисовать markdown-текст в документ Word."""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Изображение ![подпись](путь)
        m_img = IMAGE.match(stripped)
        if m_img:
            add_image(doc, m_img.group(1), m_img.group(2), base_dir)
            i += 1
            continue

        # Блок кода
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # Таблица
        if is_table_row(line) and i + 1 < n and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(line)]
            i += 2
            while i < n and is_table_row(lines[i]):
                rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # Пустая строка
        if not stripped:
            i += 1
            continue

        # Горизонтальная линия
        if stripped in ("---", "***", "___"):
            doc.add_paragraph().add_run("─" * 40)
            i += 1
            continue

        # Заголовки
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Цитата
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, stripped[1:].strip())
            i += 1
            continue

        # Нумерованный список
        m_num = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m_num.group(2))
            i += 1
            continue

        # Маркированный список
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # Обычный абзац
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1


def set_base_style(doc: Document) -> None:
    """Базовый шрифт документа."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def convert_file(md_path: Path, out_path: Path) -> None:
    """Сконвертировать один markdown-файл в .docx."""
    doc = Document()
    set_base_style(doc)
    render_markdown(doc, md_path.read_text(encoding="utf-8"), base_dir=md_path.parent)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  + {out_path.name}")


def folder_label(folder_name: str) -> tuple[str, str]:
    """'Л01_Введение' -> ('01', 'Введение')."""
    m = re.match(r"Л(\d+)_(.*)", folder_name)
    if m:
        return m.group(1), m.group(2)
    return "", folder_name


def lab_label(file_name: str) -> tuple[str, str]:
    """'Лаб01_ХХ_и_КЗ.md' -> ('01', 'ХХ_и_КЗ')."""
    m = re.match(r"Лаб(\d+)_(.*)\.md", file_name)
    if m:
        return m.group(1), m.group(2)
    return "", file_name


def first_heading(md_path: Path) -> str:
    """Первая строка-заголовок файла."""
    for line in md_path.read_text(encoding="utf-8").splitlines():
        if line.strip().startswith("#"):
            return line.lstrip("# ").strip()
    return md_path.stem


def main() -> None:
    RESULT_DIR.mkdir(exist_ok=True)
    print(f"Результаты: {RESULT_DIR}")

    # 1) Корневые документы (программа, литература, вопросы)
    for md in sorted(COURSE_DIR.glob("*.md")):
        convert_file(md, RESULT_DIR / f"{md.stem}.docx")

    # 2) Лекции по папкам Л01..Л12
    folders = sorted(p for p in COURSE_DIR.iterdir() if p.is_dir() and p.name.startswith("Л"))
    for folder in folders:
        num, title = folder_label(folder.name)
        lecture = folder / "Лекция.md"
        if lecture.exists():
            convert_file(lecture, RESULT_DIR / f"Лекция_{num}_{title}.docx")

    # 3) Лабораторные из папки Лабораторные/
    labs = sorted(LAB_DIR.glob("Лаб*.md")) if LAB_DIR.exists() else []
    for lab in labs:
        num, title = lab_label(lab.name)
        convert_file(lab, RESULT_DIR / f"Лабораторная_{num}_{title}.docx")

    # 4) Единая расширенная методичка
    build_methodichka(folders, labs)
    print("Готово.")


def build_methodichka(folders: list[Path], labs: list[Path]) -> None:
    """Собрать единую РАСШИРЕННУЮ методичку: программа + лекции + лабораторные."""
    doc = Document()
    set_base_style(doc)

    # Титульная страница
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("СИНХРОННЫЕ МАШИНЫ")
    run.bold = True
    run.font.size = Pt(28)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Расширенное методическое пособие")
    r2.font.size = Pt(16)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Бакалавриат · 12 лекций + 4 лабораторные · теория, расчёты, моделирование")
    doc.add_page_break()

    # Программа курса (если есть) — в начало методички
    program = COURSE_DIR / "00_Программа_курса.md"
    if program.exists():
        render_markdown(doc, program.read_text(encoding="utf-8"), base_dir=COURSE_DIR)
        doc.add_page_break()

    # Содержание
    doc.add_heading("Содержание", level=1)
    doc.add_heading("Лекции", level=2)
    for folder in folders:
        lecture = folder / "Лекция.md"
        if lecture.exists():
            doc.add_paragraph(first_heading(lecture), style="List Number")
    doc.add_heading("Лабораторные работы", level=2)
    for lab in labs:
        doc.add_paragraph(first_heading(lab), style="List Number")
    doc.add_page_break()

    # Раздел «Лекции»
    doc.add_heading("ЧАСТЬ I. ЛЕКЦИИ", level=1)
    doc.add_page_break()
    lec_count = 0
    for folder in folders:
        lecture = folder / "Лекция.md"
        if not lecture.exists():
            continue
        render_markdown(doc, lecture.read_text(encoding="utf-8"), base_dir=folder)
        doc.add_page_break()
        lec_count += 1

    # Раздел «Лабораторные»
    doc.add_heading("ЧАСТЬ II. ЛАБОРАТОРНЫЕ РАБОТЫ", level=1)
    doc.add_page_break()
    lab_count = 0
    for idx, lab in enumerate(labs):
        render_markdown(doc, lab.read_text(encoding="utf-8"), base_dir=lab.parent)
        if idx < len(labs) - 1:
            doc.add_page_break()
        lab_count += 1

    out = RESULT_DIR / "Методичка_Синхронные_машины.docx"
    doc.save(out)
    print(f"  * {out.name} (лекций: {lec_count}, лабораторных: {lab_count})")


if __name__ == "__main__":
    main()
